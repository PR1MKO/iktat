from pathlib import Path
from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import InvestigationAttachment
from app.paths import ensure_investigation_folder
from tests.helpers import (
    create_investigation,
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

ERTESITES_TEMPLATE_FILENAME = "ertesites_szakertoi_vizsgalatrol.docx"


def _file_safe(name: str) -> str:
    return (name or "").replace(":", "-").replace("/", "-").strip(" .")


def _expected_output_filename(case_number: str) -> str:
    safe_case = _file_safe(case_number)
    if not safe_case:
        return ERTESITES_TEMPLATE_FILENAME
    return f"{safe_case}_{ERTESITES_TEMPLATE_FILENAME}"


def _prepare_ertesites_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.subject_name = "Vizsgált Alany"
        inv.external_case_number = "EXT-42"
        inv.institution_name = "Beküldő Intézmény"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = ensure_investigation_folder(inv.case_number)
        template_dir = case_folder / "DO-NOT-EDIT"
        template_dir.mkdir(parents=True, exist_ok=True)

        from docx import Document

        template_dst = template_dir / ERTESITES_TEMPLATE_FILENAME
        doc = Document()
        doc.add_paragraph("Címzett: {{cimzett}}")
        doc.add_paragraph("Külső ügyirat: {{kulso_ugyirat}}")
        doc.add_paragraph("Ügyszám: {{iktatasi_szam}}")
        doc.add_paragraph("Jegyzőkönyv vezető: {{jkv_vezeto}}")
        doc.add_paragraph("Beküldő: {{kirendelo}}")
        doc.add_paragraph("Szak: {{szak}}")
        doc.add_paragraph("Létrehozva: {{creation_date}}")
        doc.add_paragraph("Titulus: {{titulus}}")
        doc.add_paragraph("Vizsgálat időpontja: {{vizsg_date}}")
        doc.save(template_dst)

        return {
            "inv_id": inv.id,
            "case_number": inv.case_number,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
        }


@pytest.mark.usefixtures("app")
def test_leiro_ertesites_form_generates_document(app, client, tmp_path):
    info = _prepare_ertesites_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    get_resp = client.get(f"/investigations/{info['inv_id']}/leiro/ertesites_form")
    assert get_resp.status_code == 200

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/ertesites_form",
        data={"titulus": "Dr.", "vizsg_date": "2025-09-24T10:30"},
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 303}
    assert "generated_id=" in post_resp.headers.get("Location", "")

    safe_case = _file_safe(info["case_number"])
    output_filename = _expected_output_filename(info["case_number"])
    output_path = (
        Path(app.config["INVESTIGATION_UPLOAD_FOLDER"]) / safe_case / output_filename
    )
    assert output_path.exists()

    with open(output_path, "rb") as fh:
        signature = fh.read(4)
    assert signature.startswith(b"PK")

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Vizsgált Alany" in document_xml
    assert "EXT-42" in document_xml
    assert "Beküldő Intézmény" in document_xml
    assert "Leíró Példa" in document_xml
    assert "Szakértő Példa" in document_xml
    assert "Dr." in document_xml
    assert "2025.09.24 10:30" in document_xml
    assert "{{" not in document_xml

    with app.app_context():
        expected_name = _expected_output_filename(info["case_number"])
        attachments = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"], filename=expected_name
        ).all()
        assert attachments
        att = attachments[0]
        assert att.category == "egyéb"
        assert att.uploaded_by == info["leiro_id"]


def test_leiro_sees_download_link_after_generation(app, client, tmp_path):
    info = _prepare_ertesites_case(app, tmp_path)
    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/ertesites_form",
        data={"titulus": "Dr.", "vizsg_date": "2025-09-24T10:30"},
        follow_redirects=False,
    )
    follow_resp = client.get(post_resp.headers["Location"])
    assert follow_resp.status_code == 200
    expected_name = _expected_output_filename(info["case_number"])
    assert expected_name.encode() in follow_resp.data


def test_generated_attachment_surfaces_in_documents_listing(app, client, tmp_path):
    info = _prepare_ertesites_case(app, tmp_path)
    login_follow(client, info["leiro_username"], "secret")

    client.post(
        f"/investigations/{info['inv_id']}/leiro/ertesites_form",
        data={"titulus": "Dr.", "vizsg_date": "2025-09-24T10:30"},
        follow_redirects=True,
    )

    docs_resp = client.get(f"/investigations/{info['inv_id']}/documents")
    assert docs_resp.status_code == 200
    expected_name = _expected_output_filename(info["case_number"])
    assert expected_name.encode() in docs_resp.data


def test_leiro_can_download_generated_attachment(app, client, tmp_path):
    info = _prepare_ertesites_case(app, tmp_path)
    login_follow(client, info["leiro_username"], "secret")

    client.post(
        f"/investigations/{info['inv_id']}/leiro/ertesites_form",
        data={"titulus": "Dr.", "vizsg_date": "2025-09-24T10:30"},
        follow_redirects=True,
    )

    with app.app_context():
        expected_name = _expected_output_filename(info["case_number"])
        att = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"], filename=expected_name
        ).first()
        assert att is not None
        file_id = att.id

    download_resp = client.get(f"/investigations/{info['inv_id']}/download/{file_id}")
    assert download_resp.status_code == 200
    content_disposition = download_resp.headers.get("Content-Disposition", "")
    assert content_disposition.lower().startswith("attachment;")


@pytest.mark.usefixtures("app")
def test_leiro_ertesites_form_blocks_non_leiro(app, client):
    with app.app_context():
        inv = create_investigation()
        user = create_user("not_leiro", "secret", "admin")
        inv_id = inv.id

    login_follow(client, user.username, "secret")
    post_resp = client.post(
        f"/investigations/{inv_id}/leiro/ertesites_form",
        data={"titulus": "Dr.", "vizsg_date": "2025-09-24T10:30"},
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
