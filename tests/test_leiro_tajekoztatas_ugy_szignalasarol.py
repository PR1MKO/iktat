from pathlib import Path
from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import InvestigationAttachment
from app.paths import file_safe_case_number
from tests.helpers import (
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

TEMPLATE_FILENAME = "tajekoztatas_ugy_szignalasarol.docx"


def _prepare_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    storage_root.mkdir(parents=True, exist_ok=True)

    fallback_dir = Path(app.instance_path) / "docs" / "vizsgalat"
    fallback_dir.mkdir(parents=True, exist_ok=True)

    from docx import Document

    template_path = fallback_dir / TEMPLATE_FILENAME
    doc = Document()
    doc.add_paragraph("Címzett szerv: {{cimzett-szerv}}")
    doc.add_paragraph("Titulus szerv: {{titulus-szerv}}")
    doc.add_paragraph("Kirendelő: {{kirendelo}}")
    doc.add_paragraph("Külső ügyirat: {{kulso ugyirat}}")
    doc.add_paragraph("Iktatási szám: {{iktatasi szam}}")
    doc.add_paragraph("Vezető: {{vezeto}}")
    doc.add_paragraph("Szak: {{szak}}")
    doc.add_paragraph("Actor: {{actor}}")
    doc.add_paragraph("Titulus: {{titulus}}")
    doc.add_paragraph("Dátum: {{creation_date}}")
    doc.save(template_path)

    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.external_case_number = "KULSO-42"
        inv.institution_name = "Beküldő Intézmény"
        inv.case_number = inv.case_number or "V:0001/2020"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = storage_root / file_safe_case_number(inv.case_number)
        do_not_edit = case_folder / "DO-NOT-EDIT"
        do_not_edit.mkdir(parents=True, exist_ok=True)
        template_copy = do_not_edit / TEMPLATE_FILENAME
        if template_copy.exists():
            template_copy.unlink()

        info = {
            "inv_id": inv.id,
            "case_number": inv.case_number,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
            "storage_root": storage_root,
        }

    return info


def test_get_renders_form(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_ugy_szignalasarol"
    )
    assert resp.status_code == 200
    assert "Adatok".encode("utf-8") in resp.data
    assert "Tájékoztatás ügy szignálásáról".encode("utf-8") in resp.data


def test_post_generates_document(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_ugy_szignalasarol",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "actor": "dr. Kovács Anna",
            "titulus": "intézetvezető",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code in {302, 303}
    location = post_resp.headers.get("Location", "")
    assert "generated_id=" in location

    safe_case = file_safe_case_number(info["case_number"])
    case_folder = info["storage_root"] / safe_case
    output_path = case_folder / f"{safe_case}_tajekoztatas_ugy_szignalasarol.docx"
    template_copy = case_folder / "DO-NOT-EDIT" / "tajekoztatas_ugy_szignalasarol.docx"

    assert template_copy.exists()
    assert output_path.exists()

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Országos Rendőr-főkapitányság" in document_xml
    assert "Osztályvezető" in document_xml
    assert "KULSO-42" in document_xml
    assert "Beküldő Intézmény" in document_xml
    assert info["case_number"] in document_xml
    assert "Leíró Példa" in document_xml
    assert "Szakértő Példa" in document_xml
    assert "dr. Kovács Anna" in document_xml
    assert "intézetvezető" in document_xml

    assert "Dátum: " in document_xml
    assert "{{" not in document_xml

    with app.app_context():
        attachment = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"],
            filename=output_path.name,
        ).first()
        assert attachment is not None
        assert attachment.category == "generated"
        assert attachment.uploaded_by == info["leiro_id"]

    follow_resp = client.get(location)
    assert follow_resp.status_code == 200
    assert "Letöltés".encode("utf-8") in follow_resp.data


def test_post_missing_required_fields(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_ugy_szignalasarol",
        data={
            "cimzett_szerv": "",
            "titulus_szerv": "",
            "actor": "",
            "titulus": "",
        },
        follow_redirects=False,
    )

    assert resp.status_code == 400
    assert "Hiányzó kötelező mezők".encode("utf-8") in resp.data


def test_rbac_denies_non_leiro(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    with app.app_context():
        outsider = create_user("not_leiro", "secret", "admin")

    login_follow(client, outsider.username, "secret")

    get_resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_ugy_szignalasarol"
    )
    assert get_resp.status_code in {302, 403}

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_ugy_szignalasarol",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "actor": "dr. Kovács Anna",
            "titulus": "intézetvezető",
        },
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
