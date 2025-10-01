from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import Investigation, InvestigationAttachment
from app.paths import ensure_investigation_folder, file_safe_case_number
from tests.helpers import (
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

HATARIDO_TEMPLATE_FILENAME = "hatarido_hosszabbitas_kerelem.docx"


def _prepare_hatarido_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    storage_root.mkdir(parents=True, exist_ok=True)

    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.external_case_number = "KULSO-99"
        inv.institution_name = "Beküldő Intézmény"
        inv.case_number = inv.case_number or "V:0002/2020"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = ensure_investigation_folder(inv.case_number)
        template_dir = case_folder / "DO-NOT-EDIT"
        template_dir.mkdir(parents=True, exist_ok=True)

        from docx import Document

        template_path = template_dir / HATARIDO_TEMPLATE_FILENAME
        doc = Document()
        doc.add_paragraph("Titulus szerv: {{titulus_szerv}}")
        doc.add_paragraph("Címzett szerv: {{cimzett_szerv}}")
        doc.add_paragraph("Actor: {{actor}}")
        doc.add_paragraph("Indok: {{reasons}}")
        doc.add_paragraph("Kívánt határidő: {{desired_date}}")
        doc.add_paragraph("Titulus: {{titulus}}")
        doc.add_paragraph("Iktatási szám: {{iktatasi_szam}}")
        doc.add_paragraph("Vezető: {{vezeto}}")
        doc.add_paragraph("Külső ügyirat: {{kulso_ugyirat}}")
        doc.add_paragraph("Kirendelő: {{kirendelo}}")
        doc.add_paragraph("Készítés dátuma: {{creation_date}}")
        doc.add_paragraph("Szakértő: {{szak}}")
        doc.save(template_path)

        return {
            "inv_id": inv.id,
            "case_number": inv.case_number,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
        }


@pytest.mark.usefixtures("app")
def test_leiro_hatarido_form_get(app, client, tmp_path):
    info = _prepare_hatarido_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/hatarido_hosszabbitas_kerelem"
    )
    assert resp.status_code == 200
    assert "Határidő hosszabbítás kérelem".encode("utf-8") in resp.data


@pytest.mark.usefixtures("app")
def test_leiro_hatarido_generates_document(app, client, tmp_path):
    info = _prepare_hatarido_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/hatarido_hosszabbitas_kerelem",
        data={
            "titulus_szerv": "Főosztályvezető",
            "cimzett_szerv": "Példa Szerv",
            "actor": "Dr. Valaki",
            "reasons": "túlterheltsége",
            "desired_date": "2025-12-31",
            "titulus": "Dr.",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code in {302, 303}
    location = post_resp.headers.get("Location", "")
    assert "generated_id=" in location

    follow_resp = client.get(location)
    assert follow_resp.status_code == 200
    assert "Letöltés".encode("utf-8") in follow_resp.data

    case_folder = ensure_investigation_folder(info["case_number"])
    safe_case = file_safe_case_number(info["case_number"])
    output_path = case_folder / f"{safe_case}_hatarido_hosszabbitas_kerelem.docx"
    assert output_path.exists()

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Főosztályvezető" in document_xml
    assert "Példa Szerv" in document_xml
    assert "Dr. Valaki" in document_xml
    assert "túlterheltsége" in document_xml
    assert "2025.12.31" in document_xml
    assert "Dr." in document_xml
    assert info["case_number"] in document_xml
    assert "KULSO-99" in document_xml
    assert "Beküldő Intézmény" in document_xml
    assert "Leíró Példa" in document_xml
    assert "Szakértő Példa" in document_xml

    assert "{{" not in document_xml

    with app.app_context():
        attachment = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"],
            filename=f"{safe_case}_hatarido_hosszabbitas_kerelem.docx",
        ).first()
        assert attachment is not None
        assert attachment.category == "generated"
        assert attachment.uploaded_by == info["leiro_id"]


@pytest.mark.usefixtures("app")
def test_leiro_hatarido_missing_required_fields(app, client, tmp_path):
    info = _prepare_hatarido_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/hatarido_hosszabbitas_kerelem",
        data={
            "titulus_szerv": "",
            "cimzett_szerv": "",
            "actor": "",
            "reasons": "",
            "desired_date": "",
            "titulus": "",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code == 400
    assert "Hiányzó kötelező mezők".encode("utf-8") in post_resp.data


@pytest.mark.usefixtures("app")
def test_leiro_hatarido_blocks_non_leiro(app, client, tmp_path):
    info = _prepare_hatarido_case(app, tmp_path)

    with app.app_context():
        inv = db.session.get(Investigation, info["inv_id"])
        assert inv is not None
        non_leiro = create_user("not_leiro", "secret", "admin")

    login_follow(client, non_leiro.username, "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/hatarido_hosszabbitas_kerelem"
    )
    assert resp.status_code in {302, 403}

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/hatarido_hosszabbitas_kerelem",
        data={
            "titulus_szerv": "Főosztályvezető",
            "cimzett_szerv": "Példa Szerv",
            "actor": "Dr. Valaki",
            "reasons": "túlterheltsége",
            "desired_date": "2025-12-31",
            "titulus": "Dr.",
        },
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
