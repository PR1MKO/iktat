from pathlib import Path
from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import Investigation, InvestigationAttachment
from app.paths import ensure_investigation_folder
from tests.helpers import (
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

TAJEKOZTATAS_TEMPLATE_FILENAME = "tajekoztatas_arajanlat.docx"


def _prepare_tajekoztatas_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    storage_root.mkdir(parents=True, exist_ok=True)

    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.external_case_number = "KULSO-42"
        inv.institution_name = "Beküldő Intézmény"
        inv.case_number = inv.case_number or "V:0001/2020"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = ensure_investigation_folder(inv.case_number)
        template_dir = case_folder / "DO-NOT-EDIT"
        template_dir.mkdir(parents=True, exist_ok=True)

        from docx import Document

        template_path = template_dir / TAJEKOZTATAS_TEMPLATE_FILENAME
        doc = Document()
        doc.add_paragraph("Címzett: {{cimzett-szerv}}")
        doc.add_paragraph("Titulus szerv: {{titulus-szerv}}")
        doc.add_paragraph("Külső ügyirat: {{kulso ugyirat}}")
        doc.add_paragraph("Kirendelő: {{kirendelo}}")
        doc.add_paragraph("Iktatási szám: {{iktatasi szam}}")
        doc.add_paragraph("Vezető: {{vezeto}}")
        doc.add_paragraph("Szak: {{szak}}")
        doc.add_paragraph("Actor: {{actor}}")
        doc.add_paragraph("Titulus: {{titulus}}")
        doc.add_paragraph("Összeg: {{sum}}")
        doc.add_paragraph("Dátum: {{creation_date}}")
        doc.save(template_path)

        return {
            "inv_id": inv.id,
            "case_number": inv.case_number,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
        }


@pytest.mark.usefixtures("app")
def test_leiro_tajekoztatas_form_get(app, client, tmp_path):
    info = _prepare_tajekoztatas_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.get(f"/investigations/{info['inv_id']}/leiro/tajekoztatas_arajanlat")
    assert resp.status_code == 200
    assert "Tájékoztatás árajánlat".encode("utf-8") in resp.data


@pytest.mark.usefixtures("app")
def test_leiro_tajekoztatas_generates_document(app, client, tmp_path):
    info = _prepare_tajekoztatas_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_arajanlat",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "kulso_ugyirat": "KÜ-123/2025",
            "kirendelo": "Kirendelő Intézmény",
            "actor": "N. N.",
            "titulus": "dr.",
            "sum": "250 000 Ft",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code in {302, 303}
    location = post_resp.headers.get("Location", "")
    assert "generated_id=" in location

    case_folder = ensure_investigation_folder(info["case_number"])
    output_path = Path(case_folder) / "tajekoztatas_arajanlat.docx"
    assert output_path.exists()

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Országos Rendőr-főkapitányság" in document_xml
    assert "Osztályvezető" in document_xml
    assert "KÜ-123/2025" in document_xml
    assert "Kirendelő Intézmény" in document_xml
    assert info["case_number"] in document_xml
    assert "Leíró Példa" in document_xml
    assert "Szakértő Példa" in document_xml
    assert "N. N." in document_xml
    assert "dr." in document_xml
    assert "250 000 Ft" in document_xml
    assert "{{" not in document_xml

    with app.app_context():
        attachment = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"],
            filename="tajekoztatas_arajanlat.docx",
        ).first()
        assert attachment is not None
        assert attachment.category == "generated"
        assert attachment.uploaded_by == info["leiro_id"]


@pytest.mark.usefixtures("app")
def test_leiro_tajekoztatas_missing_required_fields(app, client, tmp_path):
    info = _prepare_tajekoztatas_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_arajanlat",
        data={
            "cimzett_szerv": "",
            "titulus_szerv": "",
            "actor": "",
            "titulus": "",
            "sum": "",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code == 400
    assert "Hiányzó kötelező mezők".encode("utf-8") in post_resp.data


@pytest.mark.usefixtures("app")
def test_leiro_tajekoztatas_blocks_non_leiro(app, client, tmp_path):
    info = _prepare_tajekoztatas_case(app, tmp_path)

    with app.app_context():
        inv = db.session.get(Investigation, info["inv_id"])
        assert inv is not None
        non_leiro = create_user("not_leiro", "secret", "admin")

    login_follow(client, non_leiro.username, "secret")

    resp = client.get(f"/investigations/{info['inv_id']}/leiro/tajekoztatas_arajanlat")
    assert resp.status_code in {302, 403}

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/tajekoztatas_arajanlat",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "actor": "N. N.",
            "titulus": "dr.",
            "sum": "250 000 Ft",
        },
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
