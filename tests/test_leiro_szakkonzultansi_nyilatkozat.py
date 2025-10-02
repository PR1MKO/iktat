from pathlib import Path
from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import Investigation, InvestigationAttachment
from app.paths import ensure_investigation_folder, file_safe_case_number
from tests.helpers import (
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

TEMPLATE_FILENAME = "szakkonzultansi_nyilatkozat.docx"


def _prepare_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    storage_root.mkdir(parents=True, exist_ok=True)

    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.case_number = inv.case_number or "V:0001/2020"
        inv.external_case_number = "KULSO-42"
        inv.institution_name = "Beküldő Intézmény"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = ensure_investigation_folder(inv.case_number)

        fallback_dir = Path(app.instance_path) / "docs" / "vizsgalat"
        fallback_dir.mkdir(parents=True, exist_ok=True)

        from docx import Document

        template_path = fallback_dir / TEMPLATE_FILENAME
        doc = Document()
        doc.add_paragraph("Kirendelő: {{kirendelo}}")
        doc.add_paragraph("Külső ügyirat: {{kulso ugyirat}}")
        doc.add_paragraph("Iktatási szám: {{iktatasi szam}}")
        doc.add_paragraph("Szakértő: {{szak}}")
        doc.add_paragraph("Actor: {{actor}}")
        doc.add_paragraph("Szakterület: {{szakterulet}}")
        doc.add_paragraph("Szakkonzultáns: {{szak-cons}}")
        doc.add_paragraph("Titulus: {{titulus-cons}}")
        doc.add_paragraph("Dátum: {{creation_date}}")
        doc.save(template_path)

        return {
            "inv_id": inv.id,
            "case_number": inv.case_number,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
            "case_folder": Path(case_folder),
        }


@pytest.mark.usefixtures("app")
def test_get_renders_form(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/szakkonzultansi_nyilatkozat"
    )

    assert resp.status_code == 200
    assert "Szakkonzultánsi nyilatkozat".encode("utf-8") in resp.data


@pytest.mark.usefixtures("app")
def test_post_generates_document_and_attachment(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    payload = {
        "actor": "dr. Kovács Anna",
        "szakterulet": "Igazságügyi orvostan",
        "szak_cons": "dr. Nagy Béla",
        "titulus_cons": "szakmai konzulens",
    }

    resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakkonzultansi_nyilatkozat",
        data=payload,
        follow_redirects=False,
    )

    assert resp.status_code in {302, 303}
    location = resp.headers.get("Location", "")
    assert "generated_id=" in location

    safe_case = file_safe_case_number(info["case_number"])
    output_path = info["case_folder"] / f"{safe_case}_szakkonzultansi_nyilatkozat.docx"
    assert output_path.exists()

    template_copy = info["case_folder"] / "DO-NOT-EDIT" / TEMPLATE_FILENAME
    assert template_copy.exists()

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Beküldő Intézmény" in document_xml
    assert "KULSO-42" in document_xml
    assert "dr. Kovács Anna" in document_xml
    assert "Igazságügyi orvostan" in document_xml
    assert "dr. Nagy Béla" in document_xml
    assert "szakmai konzulens" in document_xml
    assert "Szakértő Példa" in document_xml
    assert "{{" not in document_xml

    with app.app_context():
        attachment = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"],
            filename=output_path.name,
        ).first()
        assert attachment is not None
        assert attachment.category == "generated"
        assert attachment.uploaded_by == info["leiro_id"]

    follow_resp = client.get(location)
    assert follow_resp.status_code == 200
    assert "Letöltés".encode("utf-8") in follow_resp.data


@pytest.mark.usefixtures("app")
def test_post_missing_required_fields(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakkonzultansi_nyilatkozat",
        data={
            "actor": "",
            "szakterulet": "",
            "szak_cons": "",
            "titulus_cons": "",
        },
        follow_redirects=False,
    )

    assert resp.status_code == 400
    assert "Hiányzó kötelező mezők".encode("utf-8") in resp.data


@pytest.mark.usefixtures("app")
def test_rbac_enforced(app, client, tmp_path):
    info = _prepare_case(app, tmp_path)

    with app.app_context():
        inv = db.session.get(Investigation, info["inv_id"])
        assert inv is not None
        other_user = create_user("notleiro", "secret", "admin")

    login_follow(client, other_user.username, "secret")

    get_resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/szakkonzultansi_nyilatkozat"
    )
    assert get_resp.status_code in {302, 403}

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakkonzultansi_nyilatkozat",
        data={
            "actor": "dr. Kovács Anna",
            "szakterulet": "Igazságügyi orvostan",
            "szak_cons": "dr. Nagy Béla",
            "titulus_cons": "szakmai konzulens",
        },
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
