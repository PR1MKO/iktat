import os
import shutil
import pytest

try:
    from docx import Document
except ImportError:
    pytest.skip("Codex runtime does not support 'python-docx'. Skipping tox doc tests.", allow_module_level=True)

from app.models import Case, UploadedFile, db
from tests.helpers import create_user, login


def create_case():
    case = Case(case_number='0002-2025')
    db.session.add(case)
    db.session.commit()
    return case


def test_tox_doc_generation_saves_and_registers(client, app):
    with app.app_context():
        create_user('tox', 'pw', role='toxi')
        case = create_case()
        cid = case.id

    upload_root = os.path.join(app.root_path, 'uploads')
    if os.path.exists(upload_root):
        shutil.rmtree(upload_root)

    tpl_dir = os.path.join(upload_root, 'autofill-word-do-not-edit')
    os.makedirs(tpl_dir, exist_ok=True)
    template_path = os.path.join(tpl_dir, 'Toxikológiai-kirendelő.docx')
    doc = Document()
    doc.add_paragraph('{{case.case_number}}')
    doc.add_paragraph('{{case.anyja_neve}}')
    doc.save(template_path)

    form_data = {
        'alkohol_minta_count': '1',
        'alkohol_minta_ara': '100',
    }
    with client:
        login(client, 'tox', 'pw')
        resp = client.post(f'/cases/{cid}/generate_tox_doc', data=form_data)
        assert resp.status_code == 302

    out_path = os.path.join(upload_root, case.case_number, 'Toxikológiai-kirendelő-kitöltött.docx')
    assert os.path.exists(out_path)
    assert not os.path.exists(
        os.path.join(
            upload_root,
            case.case_number,
            'webfill-do-not-edit',
            'Toxikológiai-kirendelő-kitöltött.docx',
        )
    )

    out_doc = Document(out_path)
    texts = [p.text for p in out_doc.paragraphs]
    assert case.case_number in texts
    assert 'Test Mother' in texts

    with app.app_context():
        rec = UploadedFile.query.filter_by(case_id=cid, filename='Toxikológiai-kirendelő-kitöltött.docx').first()
        assert rec is not None
        assert rec.category == 'Toxikológiai kirendelő'
