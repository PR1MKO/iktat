from pathlib import Path
from zipfile import ZipFile

import pytest

from app import db
from app.investigations.models import Investigation, InvestigationAttachment
from app.paths import ensure_investigation_folder, file_safe_case_number
from tests.helpers import (
    create_investigation_with_default_leiro,
    create_user,
    login_follow,
)

SZAKERTO_SZAKKONZULTANS_TEMPLATE_FILENAME = "szakerto-szakkonzultans_bevonasa.docx"


def _prepare_szakerto_case(app, tmp_path):
    storage_root = tmp_path / "generated-investigations"
    storage_root.mkdir(parents=True, exist_ok=True)

    with app.app_context():
        inv, leiro_user, expert_user = create_investigation_with_default_leiro()
        inv.external_case_number = "KULSO-99"
        inv.institution_name = "Beküldő Intézmény"
        db.session.commit()

        leiro_user.full_name = "Leíró Példa"
        expert_user.full_name = "Szakértő Példa"
        db.session.commit()

        app.config["INVESTIGATION_UPLOAD_FOLDER"] = str(storage_root)
        app.config["UPLOAD_INVESTIGATIONS_ROOT"] = str(storage_root)

        case_folder = ensure_investigation_folder(inv.case_number)
        template_dir = Path(case_folder) / "DO-NOT-EDIT"
        template_dir.mkdir(parents=True, exist_ok=True)

        from docx import Document

        template_path = template_dir / SZAKERTO_SZAKKONZULTANS_TEMPLATE_FILENAME
        doc = Document()
        doc.add_paragraph("Címzett: {{cimzett-szerv}}")
        doc.add_paragraph("Titulus szerv: {{titulus-szerv}}")
        doc.add_paragraph("Actor: {{actor}}")
        doc.add_paragraph("Szakterület: {{szakterulet}}")
        doc.add_paragraph("Titulus: {{titulus}}")
        doc.add_paragraph("Külső ügyirat: {{kulso ugyirat}}")
        doc.add_paragraph("Iktatási szám: {{iktatasi szam}}")
        doc.add_paragraph("Vezető: {{vezeto}}")
        doc.add_paragraph("Kirendelő: {{kirendelo}}")
        doc.add_paragraph("Szak: {{szak}}")
        doc.add_paragraph("Dátum: {{creation_date}}")
        doc.save(template_path)

        return {
            "inv_id": inv.id,
            "leiro_username": leiro_user.username,
            "leiro_id": leiro_user.id,
            "case_number": inv.case_number,
            "case_folder": Path(case_folder),
            "safe_case": file_safe_case_number(inv.case_number),
        }


@pytest.mark.usefixtures("app")
def test_leiro_szakerto_szakkonzultans_form_get(app, client, tmp_path):
    info = _prepare_szakerto_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/szakerto_szakkonzultans_bevonasa"
    )
    assert resp.status_code == 200
    assert "Szakértő / szakkonzultáns bevonása".encode("utf-8") in resp.data


@pytest.mark.usefixtures("app")
def test_leiro_szakerto_szakkonzultans_generates_document(app, client, tmp_path):
    info = _prepare_szakerto_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakerto_szakkonzultans_bevonasa",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "actor": "Dr. Valaki",
            "szakterulet": "option2",
            "titulus": "Dr.",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code in {302, 303}
    location = post_resp.headers.get("Location", "")
    assert "generated_id=" in location

    output_path = (
        Path(info["case_folder"])
        / f"{info['safe_case']}_szakerto_szakkonzultans_bevonasa.docx"
    )
    assert output_path.exists()

    with ZipFile(output_path) as bundle:
        document_xml = bundle.read("word/document.xml").decode("utf-8")

    assert "Országos Rendőr-főkapitányság" in document_xml
    assert "Osztályvezető" in document_xml
    assert "Dr. Valaki" in document_xml
    assert "option2" in document_xml
    assert "Dr." in document_xml
    assert "KULSO-99" in document_xml
    assert "Beküldő Intézmény" in document_xml
    assert "Leíró Példa" in document_xml
    assert "Szakértő Példa" in document_xml
    assert "{{" not in document_xml

    with app.app_context():
        attachment = InvestigationAttachment.query.filter_by(
            investigation_id=info["inv_id"],
            filename=output_path.name,
        ).first()
        assert attachment is not None
        assert attachment.category == "generated"
        assert attachment.uploaded_by == info["leiro_id"]

    follow_resp = client.get(location)
    assert follow_resp.status_code == 200
    assert "Letöltés".encode("utf-8") in follow_resp.data


@pytest.mark.usefixtures("app")
def test_leiro_szakerto_szakkonzultans_missing_required(app, client, tmp_path):
    info = _prepare_szakerto_case(app, tmp_path)

    login_follow(client, info["leiro_username"], "secret")

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakerto_szakkonzultans_bevonasa",
        data={
            "cimzett_szerv": "",
            "titulus_szerv": "",
            "actor": "",
            "szakterulet": "",
            "titulus": "",
        },
        follow_redirects=False,
    )

    assert post_resp.status_code == 400
    assert "Hiányzó kötelező mezők".encode("utf-8") in post_resp.data


@pytest.mark.usefixtures("app")
def test_leiro_szakerto_szakkonzultans_rbac(app, client, tmp_path):
    info = _prepare_szakerto_case(app, tmp_path)

    with app.app_context():
        inv = db.session.get(Investigation, info["inv_id"])
        assert inv is not None
        other_user = create_user("not_leiro", "secret", "admin")

    login_follow(client, other_user.username, "secret")

    resp = client.get(
        f"/investigations/{info['inv_id']}/leiro/szakerto_szakkonzultans_bevonasa"
    )
    assert resp.status_code in {302, 403}

    post_resp = client.post(
        f"/investigations/{info['inv_id']}/leiro/szakerto_szakkonzultans_bevonasa",
        data={
            "cimzett_szerv": "Országos Rendőr-főkapitányság",
            "titulus_szerv": "Osztályvezető",
            "actor": "Dr. Valaki",
            "szakterulet": "option2",
            "titulus": "Dr.",
        },
        follow_redirects=False,
    )
    assert post_resp.status_code in {302, 403}
