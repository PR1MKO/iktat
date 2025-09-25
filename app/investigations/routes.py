import io
import re
import tempfile
import unicodedata
import zipfile
from datetime import datetime, timedelta
from pathlib import Path

from flask import (
    abort,
    current_app,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    url_for,
)
from flask_login import current_user, login_required
from sqlalchemy import func, or_, text
from werkzeug import exceptions
from werkzeug.utils import secure_filename  # for safe filenames

from app import db
from app.models import User
from app.paths import ensure_investigation_folder
from app.services.case_logic import resolve_effective_describer_user
from app.services.core_user_read import get_user_safe

# IMPORTANT: import the module (so monkeypatch in tests affects calls)
from app.utils import permissions as permissions_mod
from app.utils.dates import safe_fmt
from app.utils.rbac import require_roles as roles_required
from app.utils.roles import canonical_role
from app.utils.time_utils import fmt_budapest, fmt_date, now_utc
from app.utils.uploads import send_safe

from . import investigations_bp
from .forms import FileUploadForm, InvestigationForm, InvestigationNoteForm
from .models import (
    Investigation,
    InvestigationAttachment,
    InvestigationChangeLog,
    InvestigationNote,
)
from .utils import (
    display_name,
    generate_case_number,
    init_investigation_upload_dirs,
    user_display_name,
)

# ---------------------------------------------------------------------------
# Role helpers
# ---------------------------------------------------------------------------


def normalize_role(raw):
    """Map legacy/diacritic variants to canonical short codes."""
    can = canonical_role(raw)
    if can:
        return can
    mapping = {
        # szignáló → szig
        "szignáló": "szig",
        "szignalo": "szig",
        "szignaló": "szig",
        # szakértő → szak
        "szakértő": "szak",
        "szakerto": "szak",
        # leíró → leir
        "leíró": "leir",
        "leiro": "leir",
        # pénzügy → penz
        "pénzügy": "penz",
        "penzugy": "penz",
        "pénz": "penz",
    }
    key = (raw or "").strip().casefold()
    return mapping.get(key, (raw or "").strip())


# Keep the old name but delegate to the module, so tests that monkeypatch
# app.utils.permissions.capabilities_for are respected here.
def capabilities_for(user):
    return permissions_mod.capabilities_for(user)


# Permission helpers for upload UI — accept canonical AND legacy labels
ALWAYS_UPLOAD_ROLES = {
    "admin",
    "iroda",
    "szig",
    "szignáló",  # legacy spelling
    "penz",
    "pénzügy",  # explicit diacritic alias
}
CONDITIONAL_UPLOAD_ROLES = {
    "szak",  # canonical
    "szakértő",  # legacy with accents
    "szakerto",  # legacy ascii
    "leir",  # canonical scribe
    "leíró",  # legacy with accents
    "leiro",  # ascii
    "toxi",
}


def _is_assigned_member(inv, u):
    uid = getattr(u, "id", None)
    if not uid:
        return False

    explicit_describer_id = getattr(inv, "describer_id", None)
    if uid in {
        getattr(inv, "expert1_id", None),
        getattr(inv, "expert2_id", None),
        explicit_describer_id,
    }:
        return True

    if explicit_describer_id:
        return False

    for expert_id in (
        getattr(inv, "expert1_id", None),
        getattr(inv, "expert2_id", None),
    ):
        if not expert_id:
            continue
        expert_user = db.session.get(User, expert_id)
        effective = resolve_effective_describer_user(expert_user, explicit_describer_id)
        if effective and effective.id == uid:
            return True

    return False


def can_upload_investigation_now(inv, u):
    # Normalize but also consider the raw role string to be belt-and-suspenders
    raw = getattr(u, "role", None) or ""
    role = normalize_role(raw)
    candidate = role or raw
    if candidate in ALWAYS_UPLOAD_ROLES:
        return True
    if candidate in CONDITIONAL_UPLOAD_ROLES:
        return _is_assigned_member(inv, u)
    return False


def cannot_upload_reason(inv, u):
    role = normalize_role(getattr(u, "role", None))
    if role in {
        "leir",
        "leíró",
        "leiro",
        "szak",
        "szakértő",
        "szakerto",
        "toxi",
    } and not _is_assigned_member(inv, u):
        return "Csak a kijelölt szakértő vagy leíró tölthet fel a vizsgálathoz."
    return "Nincs jogosultság a feltöltéshez."


@investigations_bp.route("/<int:id>/leiro/elvegzem", methods=["GET"])
@login_required
@roles_required("leíró", "leir", "LEIRO", "lei")
def leiro_elvegzem(id: int):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)

    form = InvestigationForm(obj=inv)

    inv.birth_date_str = fmt_date(inv.birth_date)
    inv.registration_time_str = fmt_date(inv.registration_time)
    inv.deadline_str = fmt_date(inv.deadline)

    notes = (
        InvestigationNote.query.filter_by(investigation_id=id)
        .order_by(InvestigationNote.timestamp.desc())
        .all()
    )
    attachments = (
        InvestigationAttachment.query.filter_by(investigation_id=id)
        .order_by(InvestigationAttachment.uploaded_at.desc())
        .all()
    )
    changelog = (
        InvestigationChangeLog.query.filter_by(investigation_id=id)
        .order_by(InvestigationChangeLog.timestamp.desc())
        .all()
    )

    for note in notes:
        note.timestamp_str = fmt_budapest(note.timestamp)
        note.author = get_user_safe(note.author_id)

    for att in attachments:
        att.uploaded_at_str = fmt_budapest(att.uploaded_at)

    for log in changelog:
        log.timestamp_str = fmt_budapest(log.timestamp)
        log.editor = get_user_safe(log.edited_by)

    assignment_type_label = dict(form.assignment_type.choices).get(
        inv.assignment_type, inv.assignment_type
    )
    investigation_type_label = dict(form.investigation_type.choices).get(
        inv.investigation_type, inv.investigation_type
    )

    assigned_expert_user = (
        get_user_safe(inv.assigned_expert_id) if inv.assigned_expert_id else None
    )
    assigned_expert_display = (
        user_display_name(assigned_expert_user) if assigned_expert_user else None
    )

    expert_user = getattr(inv, "expert", None) or (
        get_user_safe(getattr(inv, "expert1_id", None))
        if getattr(inv, "expert1_id", None)
        else None
    )
    if expert_user is None:
        expert_user = assigned_expert_user

    describer_user = getattr(inv, "describer", None)
    if describer_user is None and getattr(inv, "describer_id", None):
        describer_user = get_user_safe(inv.describer_id)
    if describer_user is None and expert_user is not None:
        default_leiro_id = getattr(expert_user, "default_leiro_id", None)
        if default_leiro_id:
            describer_user = get_user_safe(default_leiro_id)

    expert_display = display_name(expert_user)
    describer_display = display_name(describer_user)

    caps = dict(capabilities_for(current_user) or {})

    is_assigned = _is_assigned_member(inv, current_user)
    can_upload_ui = (
        can_upload_investigation_now(inv, current_user) if is_assigned else False
    )
    deny_reason = None if can_upload_ui else cannot_upload_reason(inv, current_user)
    caps["can_post_investigation_notes"] = bool(is_assigned)

    return render_template(
        "investigations/leiro_elvegzem.html",
        investigation=inv,
        notes=notes,
        attachments=attachments,
        changelog=changelog,
        user_display_name=user_display_name,
        caps=caps,
        can_upload_ui=can_upload_ui,
        upload_form=FileUploadForm(),
        cannot_upload_reason=deny_reason,
        assignment_type_label=assignment_type_label,
        investigation_type_label=investigation_type_label,
        assigned_expert_display=assigned_expert_display,
        expert_display_name=expert_display,
        describer_display_name=describer_display,
    )


ERTESITES_TEMPLATE_FILENAME = "ertesites_szakertoi_vizsgalatrol.docx"
ERTESITES_TEMPLATE_DIRNAME = "DO-NOT-EDIT"


@investigations_bp.route("/<int:id>/leiro/ertesites_form", methods=["GET", "POST"])
@login_required
@roles_required("leíró", "leir", "LEIRO", "lei")
def leiro_ertesites_form(id: int):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)

    form_data = {
        "titulus": request.form.get("titulus", ""),
        "vizsg_date": request.form.get("vizsg_date", ""),
    }

    if request.method == "POST":
        errors: list[str] = []

        titulus = form_data["titulus"].strip()
        if not titulus:
            errors.append("A titulus mező kitöltése kötelező.")

        vizsg_dt = None
        vizsg_raw = form_data["vizsg_date"].strip()
        if vizsg_raw:
            try:
                vizsg_dt = datetime.strptime(vizsg_raw, "%Y-%m-%dT%H:%M")
            except ValueError:
                errors.append("Érvénytelen dátum és időpont formátum.")
        else:
            errors.append("A dátum és időpont mező kitöltése kötelező.")

        if errors:
            for message in errors:
                flash(message, "danger")
            return (
                render_template(
                    "investigations/ertesites_doc_form.html",
                    investigation=inv,
                    form_data=form_data,
                ),
                400,
            )

        assigned_expert_user = (
            get_user_safe(inv.assigned_expert_id) if inv.assigned_expert_id else None
        )
        expert_user = getattr(inv, "expert", None)
        if expert_user is None:
            for candidate_id in (
                getattr(inv, "expert1_id", None),
                getattr(inv, "expert2_id", None),
            ):
                if not candidate_id:
                    continue
                expert_user = get_user_safe(candidate_id)
                if expert_user:
                    break
        if expert_user is None:
            expert_user = assigned_expert_user

        describer_user = getattr(
            inv, "describer", None
        ) or resolve_effective_describer_user(
            expert_user, getattr(inv, "describer_id", None)
        )

        creation_date = fmt_budapest(now_utc(), "%Y.%m.%d")
        vizsg_formatted = vizsg_dt.strftime("%Y.%m.%d %H:%M") if vizsg_dt else ""

        context = {
            "cimzett": (inv.subject_name or ""),
            "kulso ugyirat": (inv.external_case_number or ""),
            "iktatasi szam": (inv.case_number or ""),
            "vezeto": user_display_name(describer_user) if describer_user else "",
            "kirendelo": (inv.institution_name or ""),
            "szak": user_display_name(expert_user) if expert_user else "",
            "creation_date": creation_date,
            "titulus": titulus,
            "vizsg_date": vizsg_formatted,
        }
        context.update(
            {
                "kulso_ugyirat": context["kulso ugyirat"],
                "iktatasi_szam": context["iktatasi szam"],
                "jkv.vezető": context["vezeto"],
                "jkv_vezeto": context["vezeto"],
            }
        )

        case_folder = ensure_investigation_folder(inv.case_number)
        template_path = (
            Path(case_folder) / ERTESITES_TEMPLATE_DIRNAME / ERTESITES_TEMPLATE_FILENAME
        )
        if not template_path.exists():
            abort(404, description="A sablon nem található ehhez a vizsgálathoz.")

        output_path = Path(case_folder) / ERTESITES_TEMPLATE_FILENAME

        try:
            _render_docx_template(template_path, output_path, context)
        except Exception as exc:  # noqa: BLE001
            current_app.logger.exception(
                "DOCX generation failed for investigation %s: %s", inv.id, exc
            )
            flash("Hiba történt a dokumentum generálása közben.", "danger")
            return (
                render_template(
                    "investigations/ertesites_doc_form.html",
                    investigation=inv,
                    form_data=form_data,
                ),
                500,
            )

        flash("Értesítés dokumentum sikeresen generálva.", "success")
        return redirect(url_for("investigations.leiro_ertesites_form", id=inv.id))

    return render_template(
        "investigations/ertesites_doc_form.html",
        investigation=inv,
        form_data=form_data,
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _render_docx_template(
    template_path: Path, output_path: Path, context: dict
) -> None:
    """Render a DOCX template after sanitizing malformed Jinja placeholders."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    sanitized_path = _sanitize_docx_placeholders(template_path)
    safe_context = {k: ("" if v is None else v) for k, v in context.items()}

    try:
        try:
            from docxtpl import DocxTemplate
        except ModuleNotFoundError:
            DocxTemplate = None
        else:
            tpl = DocxTemplate(str(sanitized_path))
            tpl.render(safe_context)
            tpl.save(str(output_path))
            return

        from docx import Document

        doc = Document(str(sanitized_path))
        replacements = {
            f"{{{{{key}}}}}": (str(value) if value is not None else "")
            for key, value in safe_context.items()
        }

        for paragraph in doc.paragraphs:
            for needle, replacement in replacements.items():
                if needle in paragraph.text:
                    paragraph.text = paragraph.text.replace(needle, replacement)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    for needle, replacement in replacements.items():
                        if needle in text:
                            text = text.replace(needle, replacement)
                    if cell.text != text:
                        cell.text = text

        doc.save(str(output_path))
    finally:
        try:
            sanitized_path.unlink()
        except FileNotFoundError:
            pass


_W_T_OPEN = re.compile(rb"</w:t>\s*<w:t[^>]*>")
_PLACEHOLDER_PATTERN = re.compile(r"(\{\{.*?\}\}|\{%.*?%\}|\{#.*?#\})", flags=re.DOTALL)
_VAR_NAME_PATTERN = re.compile(r"\{\{\s*(.*?)\s*\}\}")


def _sanitize_docx_placeholders(docx_path: Path) -> Path:
    """Return sanitized copy of DOCX with normalized Jinja placeholders."""
    changed_any = False
    out_buf = io.BytesIO()

    with (
        zipfile.ZipFile(str(docx_path), "r") as zin,
        zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout,
    ):
        names = zin.namelist()
        target_parts = [
            n
            for n in names
            if n.startswith("word/")
            and n.endswith(".xml")
            and (
                n == "word/document.xml"
                or n.startswith("word/header")
                or n.startswith("word/footer")
            )
        ]
        for name in names:
            data = zin.read(name)
            if name in target_parts:
                try:
                    new_data, changed = _normalize_xml_placeholders(data)
                except Exception:
                    new_data, changed = data, False
                else:
                    data = new_data
                changed_any = changed_any or changed
            zout.writestr(name, data)

    tmp = tempfile.NamedTemporaryFile(prefix="sanitized_", suffix=".docx", delete=False)
    with tmp:
        tmp.write(out_buf.getvalue())
    try:
        current_app.logger.info(
            "DOCX sanitizer: wrote %s (changed=%s)",
            tmp.name,
            "yes" if changed_any else "no",
        )
    except Exception:
        pass
    return Path(tmp.name)


def _normalize_xml_placeholders(xml_bytes: bytes) -> tuple[bytes, bool]:
    """Normalize placeholders in a DOCX XML part."""

    joined = _W_T_OPEN.sub(b"", xml_bytes)
    changed = joined != xml_bytes
    text = joined.decode("utf-8", errors="ignore")

    def _ascii_fold(value: str) -> str:
        nfkd = unicodedata.normalize("NFKD", value)
        return "".join(ch for ch in nfkd if not unicodedata.combining(ch))

    def repl(match: re.Match[str]) -> str:
        nonlocal changed
        token = match.group(0)
        vm = _VAR_NAME_PATTERN.fullmatch(token)
        if not vm:
            return token
        var = vm.group(1)
        folded = _ascii_fold(var)
        normalized = re.sub(r"[^\w]", "_", folded)
        normalized = re.sub(r"_+", "_", normalized).strip("_")
        if normalized != var:
            changed = True
            return "{{" + normalized + "}}"
        return token

    new_text = _PLACEHOLDER_PATTERN.sub(repl, text)
    return new_text.encode("utf-8"), changed


def _can_modify(inv, user) -> bool:
    return normalize_role(user.role) in {"admin", "iroda"}


def _can_note_or_upload(inv, user) -> bool:
    r = normalize_role(user.role)
    if r in {"admin", "iroda"}:
        return True
    return user.id in {inv.expert1_id, inv.expert2_id, inv.describer_id}


def _log_changes(inv: Investigation, form: InvestigationForm):
    fields = [
        "subject_name",
        "maiden_name",
        "mother_name",
        "birth_place",
        "birth_date",
        "taj_number",
        "residence",
        "citizenship",
        "institution_name",
        "investigation_type",
        "external_case_number",
        "other_identifier",
    ]
    logs = []
    for field in fields:
        new_val = getattr(form, field).data
        old_val = getattr(inv, field)
        if old_val != new_val:
            setattr(inv, field, new_val)
            logs.append(
                InvestigationChangeLog(
                    investigation_id=inv.id,
                    field_name=field,
                    old_value=str(old_val) if old_val is not None else None,
                    new_value=str(new_val) if new_val is not None else None,
                    edited_by=current_user.id,
                    timestamp=now_utc(),
                )
            )
    return logs


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@investigations_bp.route("/")
@login_required  # Any authenticated user may access the investigations index
def list_investigations():
    search = (request.args.get("search") or request.args.get("q") or "").strip()
    case_type = request.args.get("case_type", "").strip()
    sort_by = request.args.get("sort_by", "case_number")
    sort_order = request.args.get("sort_order", "asc")
    page = request.args.get("page", 1, type=int)
    per_page = 25

    query = Investigation.query

    if search:
        like = f"%{search}%"
        query = query.filter(
            or_(
                Investigation.case_number.ilike(like),
                Investigation.external_case_number.ilike(like),
                Investigation.other_identifier.ilike(like),
                Investigation.taj_number.ilike(like),
                Investigation.subject_name.ilike(like),
                Investigation.maiden_name.ilike(like),
                Investigation.mother_name.ilike(like),
                Investigation.birth_place.ilike(like),
                Investigation.residence.ilike(like),
                Investigation.citizenship.ilike(like),
                Investigation.investigation_type.ilike(like),
                Investigation.institution_name.ilike(like),
                func.strftime("%Y-%m-%d", Investigation.birth_date).like(like),
                func.strftime("%Y", Investigation.birth_date).like(like),
            )
        )

    if case_type:
        query = query.filter(Investigation.investigation_type == case_type)

    order_col = {
        "case_number": Investigation.case_number,
        "deadline": Investigation.deadline,
    }.get(sort_by, Investigation.id)

    query = query.order_by(
        order_col.desc() if sort_order == "desc" else order_col.asc()
    )

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    investigations = pagination.items

    for inv in investigations:
        inv.registration_time_str = fmt_date(inv.registration_time)
        inv.deadline_str = fmt_date(inv.deadline)
        inv.expert1_name = user_display_name(get_user_safe(inv.expert1_id))
        inv.expert2_name = user_display_name(get_user_safe(inv.expert2_id))
        inv.describer_name = user_display_name(get_user_safe(inv.describer_id))

    has_edit_investigation = (
        "investigations.edit_investigation" in current_app.view_functions
    )

    return render_template(
        "investigations/list.html",
        investigations=investigations,
        pagination=pagination,
        sort_by=sort_by,
        sort_order=sort_order,
        case_type_filter=case_type,
        search_query=search,
        query_params=request.args.to_dict(),
        has_edit_investigation=has_edit_investigation,
        caps=capabilities_for(current_user),
    )


@investigations_bp.route("/new", methods=["GET", "POST"])
@login_required
@roles_required("admin", "iroda")
def new_investigation():
    form = InvestigationForm()

    # Include canonical and legacy labels so tests using "szakértő" are covered
    experts = db.session.execute(
        text(
            "SELECT id, screen_name, username FROM user "
            "WHERE role IN ('szak', 'szakértő', 'szakerto') "
            "ORDER BY screen_name, username"
        )
    ).all()
    form.assigned_expert_id.choices = [(0, "— Válasszon —")] + [
        (row.id, row.screen_name or row.username) for row in experts
    ]

    if form.validate_on_submit():
        assignment_type = form.assignment_type.data
        assigned_expert_id = (
            form.assigned_expert_id.data if assignment_type == "SZAKÉRTŐI" else None
        )
        inv = Investigation(
            subject_name=form.subject_name.data,
            maiden_name=(
                getattr(form, "maiden_name", None).data
                if hasattr(form, "maiden_name")
                else None
            ),
            mother_name=form.mother_name.data,
            birth_place=form.birth_place.data,
            birth_date=form.birth_date.data,
            taj_number=form.taj_number.data,
            residence=form.residence.data,
            citizenship=form.citizenship.data,
            institution_name=form.institution_name.data,
            investigation_type=form.investigation_type.data,
            external_case_number=form.external_case_number.data,
            other_identifier=form.other_identifier.data,
            assignment_type=assignment_type,
            assigned_expert_id=assigned_expert_id,
            status="beérkezett",
        )
        inv.case_number = generate_case_number(db.session)  # V-####-YYYY
        inv.registration_time = now_utc()
        inv.deadline = inv.registration_time + timedelta(days=30)

        change_log_rows = []
        if assignment_type == "SZAKÉRTŐI" and assigned_expert_id is not None:
            selected_expert_id = assigned_expert_id
            inv.assigned_expert_id = selected_expert_id
            inv.expert1_id = selected_expert_id
            previous_status = inv.status
            inv.status = "szignálva"
            expert_user = get_user_safe(selected_expert_id)
            expert_display = user_display_name(expert_user) or str(selected_expert_id)
            change_log_rows.extend(
                [
                    ("expert1_id", None, expert_display),
                    ("assigned_expert_id", None, expert_display),
                    ("status", previous_status, "szignálva"),
                ]
            )

        db.session.add(inv)
        db.session.flush()

        if change_log_rows:
            timestamp = now_utc()
            logs = [
                InvestigationChangeLog(
                    investigation_id=inv.id,
                    field_name=field,
                    old_value=old_val,
                    new_value=new_val,
                    edited_by=current_user.id,
                    timestamp=timestamp,
                )
                for field, old_val, new_val in change_log_rows
            ]
            db.session.add_all(logs)

        db.session.commit()

        # Create per-investigation folder (separate from Cases)
        ensure_investigation_folder(inv.case_number)
        init_investigation_upload_dirs(inv.case_number)

        flash("Vizsgálat létrehozva.", "success")
        return redirect(url_for("investigations.documents", id=inv.id))
    elif request.method == "POST":
        for errs in form.errors.values():
            for err in errs:
                flash(err)
        if current_app.config.get("STRICT_PRG_ENABLED", True):
            return redirect(url_for("investigations.new_investigation"))

    return render_template("investigations/new.html", form=form)


@investigations_bp.route("/<int:id>/documents", methods=["GET"])
@login_required
@roles_required(
    "admin",
    "iroda",
    "szak",
    "szakértő",  # allow legacy expert label
    "leir",
    "leíró",  # allow legacy scribe label
    "szig",
    "toxi",
    "penz",
    "pénzügy",  # allow legacy finance label
)
def documents(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)

    if request.args.get("uploaded"):
        flash("Fájl feltöltve", "success")

    attachments = (
        InvestigationAttachment.query.filter_by(investigation_id=id)
        .order_by(InvestigationAttachment.uploaded_at.desc())
        .all()
    )
    for att in attachments:
        att.uploaded_at_str = safe_fmt(att.uploaded_at)

    upload_form = FileUploadForm()

    can_upload_ui = can_upload_investigation_now(inv, current_user)
    deny_reason = None if can_upload_ui else cannot_upload_reason(inv, current_user)

    return render_template(
        "investigations/documents.html",
        investigation=inv,
        attachments=attachments,
        upload_form=upload_form,
        upload_url=url_for("investigations.upload_investigation_file", id=inv.id),
        can_upload_ui=can_upload_ui,
        cannot_upload_reason=deny_reason,
    )


@investigations_bp.route("/<int:id>/view")
@login_required
@roles_required("admin", "iroda", "szak", "penz", "pénzügy", "szig")
def view_investigation(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)

    attachments = (
        InvestigationAttachment.query.filter_by(investigation_id=id)
        .order_by(InvestigationAttachment.uploaded_at.desc())
        .all()
    )

    notes = (
        InvestigationNote.query.filter_by(investigation_id=id)
        .order_by(InvestigationNote.timestamp.desc())
        .all()
    )
    for note in notes:
        note.author = get_user_safe(note.author_id)
        note.timestamp_str = safe_fmt(note.timestamp)

    changelog_entries = (
        InvestigationChangeLog.query.filter_by(investigation_id=id)
        .order_by(InvestigationChangeLog.timestamp.desc())
        .all()
    )
    for entry in changelog_entries:
        entry.edited_by = user_display_name(get_user_safe(entry.edited_by))
        entry.timestamp_str = safe_fmt(entry.timestamp)

    assigned_expert = get_user_safe(inv.assigned_expert_id)
    inv.registration_time_str = safe_fmt(inv.registration_time)
    inv.deadline_str = safe_fmt(inv.deadline)

    return render_template(
        "investigations/view.html",
        investigation=inv,
        attachments=attachments,
        notes=notes,
        assigned_expert=assigned_expert,
        changelog_entries=changelog_entries,
        user_display_name=user_display_name,
        caps=capabilities_for(current_user),
    )


@investigations_bp.route("/<int:id>")
@login_required
@roles_required(
    "admin",
    "iroda",
    "szak",
    "szakértő",  # allow legacy expert label
    "leir",
    "leíró",  # allow legacy scribe label
    "szig",
    "toxi",
    "penz",
    "pénzügy",  # allow legacy finance label
)
def detail_investigation(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)
    form = InvestigationForm(obj=inv)
    note_form = InvestigationNoteForm()
    upload_form = FileUploadForm()
    inv.birth_date_str = fmt_date(inv.birth_date)
    inv.registration_time_str = fmt_date(inv.registration_time)
    inv.deadline_str = fmt_date(inv.deadline)
    notes = (
        InvestigationNote.query.filter_by(investigation_id=id)
        .order_by(InvestigationNote.timestamp.desc())
        .all()
    )
    attachments = (
        InvestigationAttachment.query.filter_by(investigation_id=id)
        .order_by(InvestigationAttachment.uploaded_at.desc())
        .all()
    )
    for note in notes:
        note.timestamp_str = fmt_budapest(note.timestamp)
        note.author = get_user_safe(note.author_id)
    for att in attachments:
        att.uploaded_at_str = fmt_budapest(att.uploaded_at)
    changelog = (
        InvestigationChangeLog.query.filter_by(investigation_id=id)
        .order_by(InvestigationChangeLog.timestamp.desc())
        .all()
    )
    for log in changelog:
        log.timestamp_str = fmt_budapest(log.timestamp)
        log.editor = get_user_safe(log.edited_by)
    assignment_type_label = dict(form.assignment_type.choices).get(
        inv.assignment_type, inv.assignment_type
    )
    investigation_type_label = dict(form.investigation_type.choices).get(
        inv.investigation_type, inv.investigation_type
    )
    assigned_expert_user = (
        get_user_safe(inv.assigned_expert_id) if inv.assigned_expert_id else None
    )
    assigned_expert_display = (
        user_display_name(assigned_expert_user) if assigned_expert_user else None
    )

    expert_user = getattr(inv, "expert", None) or (
        get_user_safe(getattr(inv, "expert1_id", None))
        if getattr(inv, "expert1_id", None)
        else None
    )
    if expert_user is None:
        expert_user = assigned_expert_user

    describer_user = getattr(inv, "describer", None)
    if describer_user is None and getattr(inv, "describer_id", None):
        describer_user = get_user_safe(inv.describer_id)
    if describer_user is None and expert_user is not None:
        default_leiro_id = getattr(expert_user, "default_leiro_id", None)
        if default_leiro_id:
            describer_user = get_user_safe(default_leiro_id)

    expert_display_name = display_name(expert_user)
    describer_display_name = display_name(describer_user)

    caps = capabilities_for(current_user)
    can_upload_ui = can_upload_investigation_now(inv, current_user)
    deny_reason = None if can_upload_ui else cannot_upload_reason(inv, current_user)

    return render_template(
        "investigations/detail.html",
        investigation=inv,
        form=form,
        note_form=note_form,
        upload_form=upload_form,
        notes=notes,
        attachments=attachments,
        changelog=changelog,
        user_display_name=user_display_name,
        caps=caps,
        can_upload_ui=can_upload_ui,
        cannot_upload_reason=deny_reason,
        assignment_type_label=assignment_type_label,
        investigation_type_label=investigation_type_label,
        assigned_expert_display=assigned_expert_display,
        expert_display_name=expert_display_name,
        describer_display_name=describer_display_name,
    )


@investigations_bp.route("/<int:id>/edit", methods=["POST"])
@login_required
@roles_required("admin", "iroda")
def edit_investigation(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)
    caps = capabilities_for(current_user)
    if not caps.get("can_edit_investigation"):
        flash("Nincs jogosultság", "danger")
        return redirect(url_for("investigations.detail_investigation", id=id))
    form = InvestigationForm()
    if not form.validate_on_submit():
        flash("Hibás űrlap", "error")
        return redirect(url_for("investigations.detail_investigation", id=id))
    logs = _log_changes(inv, form)
    db.session.add_all(logs)
    db.session.commit()
    flash("Vizsgálat frissítve.", "success")
    return redirect(url_for("investigations.detail_investigation", id=id))


@investigations_bp.route("/<int:id>/notes", methods=["POST"])
@login_required
@roles_required("admin", "iroda", "szak", "szig", "leíró", "leir", "LEIRO", "lei")
def add_investigation_note(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)
    caps = capabilities_for(current_user) or {}
    if not caps.get("can_post_investigation_notes"):
        abort(403)
    if normalize_role(current_user.role) not in {
        "admin",
        "iroda",
    } and not _is_assigned_member(inv, current_user):
        abort(403)

    form = InvestigationNoteForm()
    text = None
    if form.validate_on_submit():
        text = form.text.data
    else:
        data = request.get_json(silent=True) or {}
        text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Empty note"}), 400

    note = InvestigationNote(
        investigation_id=inv.id, author_id=current_user.id, text=text
    )
    db.session.add(note)
    db.session.commit()

    # ensure author is available for the partial
    author = get_user_safe(note.author_id)
    note.timestamp_str = fmt_budapest(note.timestamp)

    html = render_template(
        "investigations/_note.html",
        note=note,
        author=author,
        user_display_name=user_display_name,
    )
    return html


@investigations_bp.route("/<int:id>/upload", methods=["POST"])
@login_required
@roles_required(
    "admin",
    "iroda",
    "szak",
    "szakértő",  # legacy expert label
    "leir",
    "leíró",  # legacy scribe label
    "szig",
    "szignáló",
    "toxi",
    "penz",
    "pénzügy",  # legacy finance label
)
def upload_investigation_file(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)
    if not can_upload_investigation_now(inv, current_user):
        abort(403)

    # Support CSRF-free AJAX uploads (test path) and normal multipart posts
    is_xhr = request.headers.get("X-Requested-With") == "XMLHttpRequest"
    category = (request.form.get("category") or "").strip()
    upfile = request.files.get("file")

    if not category or not upfile or not getattr(upfile, "filename", "").strip():
        if is_xhr:
            return jsonify({"error": "invalid"}), 400
        flash("Kategória megadása kötelező.", "danger")
        return redirect(url_for("investigations.documents", id=id))

    # Secondary guard on filename
    if not getattr(upfile, "filename", "").strip():
        if is_xhr:
            return jsonify({"error": "invalid"}), 400
        flash("Nincs kiválasztott fájl.", "warning")
        return redirect(url_for("investigations.documents", id=id))

    # ---- Save directly into the investigation folder (matches tests' expectation) ----
    try:
        inv_dir = Path(ensure_investigation_folder(inv.case_number))
        inv_dir.mkdir(parents=True, exist_ok=True)
        filename = secure_filename(upfile.filename or "upload.bin")
        dest = inv_dir / filename
        try:
            upfile.stream.seek(0)
        except Exception:
            pass
        upfile.save(dest.as_posix())
    except exceptions.BadRequest:
        if is_xhr:
            return jsonify({"error": "forbidden"}), 400
        flash("Érvénytelen fájl vagy kiterjesztés.", "danger")
        return redirect(url_for("investigations.documents", id=id))
    except Exception as e:
        current_app.logger.exception("Investigation file save failed: %s", e)
        if is_xhr:
            return jsonify({"error": "save-failed"}), 500
        flash("Feltöltés sikertelen.", "danger")
        return redirect(url_for("investigations.documents", id=id))

    attachment = InvestigationAttachment(
        investigation_id=inv.id,
        filename=filename,
        category=category,
        uploaded_by=current_user.id,
        uploaded_at=now_utc(),
    )
    db.session.add(attachment)
    db.session.commit()

    if is_xhr:
        return jsonify(
            {
                "id": attachment.id,
                "filename": attachment.filename,
                "category": attachment.category,
                "uploaded_at": fmt_date(attachment.uploaded_at),
            }
        )
    flash("Fájl feltöltve", "success")
    return redirect(url_for("investigations.documents", id=id, uploaded=1))


@investigations_bp.route("/<int:inv_id>/download/<int:file_id>")
@login_required
@roles_required(
    "admin",
    "iroda",
    "szak",
    "szakértő",  # allow legacy expert label
    "leir",
    "leíró",  # allow legacy scribe label
    "szig",
    "toxi",
    "penz",
    "pénzügy",  # allow legacy finance label
)
def download_investigation_file(inv_id, file_id):
    inv = db.session.get(Investigation, inv_id) or abort(404)
    att = db.session.get(InvestigationAttachment, file_id) or abort(404)
    if att.investigation_id != inv_id:
        abort(404)

    root = Path(ensure_investigation_folder(inv.case_number))
    filename = att.filename
    try:
        return send_safe(root, filename, as_attachment=True)
    except exceptions.BadRequest:
        current_app.logger.warning(
            "Path traversal attempt for investigation %s: %s", inv_id, filename
        )
        abort(400)
    except FileNotFoundError:
        current_app.logger.warning(
            "File not found for investigation %s: %s", inv_id, filename
        )
        abort(404)


# ---------------------------------------------------------------------------
# Assignment view (szignáló expert allocation)
# ---------------------------------------------------------------------------


@investigations_bp.route("/<int:id>/assign", methods=["GET", "POST"])
@login_required
@roles_required("szig")
def assign_investigation_expert(id):
    inv = db.session.get(Investigation, id)
    if inv is None:
        abort(404)

    caps = capabilities_for(current_user)
    if not caps.get("can_assign"):
        flash("Nincs jogosultság", "danger")
        return redirect(url_for("investigations.detail_investigation", id=id))

    form = InvestigationForm(obj=inv)
    note_form = InvestigationNoteForm()
    upload_form = FileUploadForm()

    inv.birth_date_str = fmt_date(inv.birth_date)
    inv.registration_time_str = fmt_date(inv.registration_time)
    inv.deadline_str = fmt_date(inv.deadline)

    notes = (
        InvestigationNote.query.filter_by(investigation_id=id)
        .order_by(InvestigationNote.timestamp.desc())
        .all()
    )
    for note in notes:
        note.timestamp_str = fmt_budapest(note.timestamp)
        note.author = get_user_safe(note.author_id)

    attachments = (
        InvestigationAttachment.query.filter_by(investigation_id=id)
        .order_by(InvestigationAttachment.uploaded_at.desc())
        .all()
    )
    for att in attachments:
        att.uploaded_at_str = fmt_budapest(att.uploaded_at)

    changelog_entries = (
        InvestigationChangeLog.query.filter_by(investigation_id=id)
        .order_by(InvestigationChangeLog.timestamp.desc())
        .all()
    )
    for entry in changelog_entries:
        entry.timestamp_str = fmt_budapest(entry.timestamp)
        entry.editor = get_user_safe(entry.edited_by)

    assignment_type_label = dict(form.assignment_type.choices).get(
        inv.assignment_type, inv.assignment_type
    )
    investigation_type_label = dict(form.investigation_type.choices).get(
        inv.investigation_type, inv.investigation_type
    )
    assigned_expert_display = None
    if inv.assigned_expert_id:
        assigned_expert_display = user_display_name(
            get_user_safe(inv.assigned_expert_id)
        )

    # Accept both canonical and legacy role labels for experts so existing
    # accounts with "szakértő" (and ASCII variants) remain visible.
    expert_role_codes = ("szak", "szakértő", "szakerto")
    szakerto_users = (
        User.query.filter(User.role.in_(expert_role_codes))
        .order_by(User.screen_name, User.username)
        .all()
    )

    def _user_label(user):
        return user.screen_name or user.username or ""

    option_items = sorted(
        [(str(u.id), _user_label(u)) for u in szakerto_users],
        key=lambda item: item[1].casefold(),
    )
    option_map = {str(u.id): u for u in szakerto_users}

    assigned_expert_display = None
    if inv.assigned_expert_id:
        expert_user = get_user_safe(inv.assigned_expert_id) or option_map.get(
            str(inv.assigned_expert_id)
        )
        assigned_expert_display = user_display_name(expert_user)

    if request.method == "POST":
        expert1_selected = (request.form.get("expert_1") or "").strip()
        expert2_selected = (request.form.get("expert_2") or "").strip()
    else:
        expert1_selected = str(inv.expert1_id) if inv.expert1_id else ""
        expert2_selected = str(inv.expert2_id) if inv.expert2_id else ""

    szakerto_choices = [("", "-- Válasszon --")] + option_items
    szakerto_choices_2 = [("", "-- Válasszon (opcionális)")] + [
        opt for opt in option_items if opt[0] != expert1_selected
    ]

    latest_log = changelog_entries[0] if changelog_entries else None
    form_version = ""
    if latest_log and latest_log.timestamp:
        form_version = latest_log.timestamp.isoformat()
    elif inv.registration_time:
        form_version = inv.registration_time.isoformat()

    # ---- UI capability for upload controls on this page:
    can_upload_ui = bool(caps.get("can_upload_investigation")) and (
        can_upload_investigation_now(inv, current_user)
    )
    deny_reason = None if can_upload_ui else cannot_upload_reason(inv, current_user)

    if request.method == "POST" and request.form.get("action") == "assign":
        expert1_raw = expert1_selected
        expert2_raw = expert2_selected

        if not expert1_raw:
            flash("Szakértő 1 kitöltése kötelező.", "warning")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        if expert1_raw not in option_map:
            flash("Érvénytelen szakértő választás.", "danger")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        if expert2_raw and expert2_raw not in option_map:
            flash("Érvénytelen második szakértő.", "danger")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        if expert2_raw and expert2_raw == expert1_raw:
            flash("A két szakértő nem lehet azonos.", "warning")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        submitted_version = (request.form.get("form_version") or "").strip()
        current_version = form_version
        if submitted_version and submitted_version != current_version:
            flash("Az űrlap időközben frissült. Kérjük, töltse be újra.", "warning")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        expert1_id = int(expert1_raw)
        expert2_id = int(expert2_raw) if expert2_raw else None

        old_expert1 = inv.expert1_id
        old_expert2 = inv.expert2_id
        old_assigned = inv.assigned_expert_id
        old_status = getattr(inv, "status", None)

        if (
            old_expert1 == expert1_id
            and (old_expert2 or None) == expert2_id
            and old_assigned == expert1_id
            and (old_status or "") == "szignálva"
        ):
            flash("Nincs változás.", "info")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        inv.expert1_id = expert1_id
        inv.expert2_id = expert2_id
        inv.assigned_expert_id = expert1_id
        if hasattr(inv, "status"):
            inv.status = "szignálva"

        timestamp = now_utc()

        def _log_change(field_name, old_val, new_val):
            if old_val == new_val:
                return None

            def _display(uid):
                if uid is None:
                    return None
                user = get_user_safe(uid) or option_map.get(str(uid))
                return user_display_name(user) if user else str(uid)

            return InvestigationChangeLog(
                investigation_id=inv.id,
                field_name=field_name,
                old_value=_display(old_val),
                new_value=_display(new_val),
                edited_by=current_user.id,
                timestamp=timestamp,
            )

        def _log_status_change(old_val, new_val):
            if old_val == new_val:
                return None
            return InvestigationChangeLog(
                investigation_id=inv.id,
                field_name="status",
                old_value=old_val,
                new_value=new_val,
                edited_by=current_user.id,
                timestamp=timestamp,
            )

        logs = list(
            filter(
                None,
                [
                    _log_change("expert1_id", old_expert1, expert1_id),
                    _log_change("expert2_id", old_expert2, expert2_id),
                    _log_change("assigned_expert_id", old_assigned, expert1_id),
                    _log_status_change(old_status, getattr(inv, "status", None)),
                ],
            )
        )

        if logs:
            db.session.add_all(logs)

        try:
            db.session.commit()
        except Exception:  # noqa: BLE001
            db.session.rollback()
            flash("Nem sikerült kijelölni a szakértőt.", "danger")
            return redirect(
                url_for("investigations.assign_investigation_expert", id=id)
            )

        flash("Szakértő kijelölve.", "success")
        return redirect(url_for("auth.szignal_cases"))

    return render_template(
        "assign_investigation_expert.html",
        investigation=inv,
        attachments=attachments,
        notes=notes,
        changelog_entries=changelog_entries,
        form=form,
        note_form=note_form,
        upload_form=upload_form,
        assignment_type_label=assignment_type_label,
        investigation_type_label=investigation_type_label,
        assigned_expert_display=assigned_expert_display,
        szakerto_choices=szakerto_choices,
        szakerto_choices_2=szakerto_choices_2,
        expert1_selected=expert1_selected,
        expert2_selected=expert2_selected,
        form_version=form_version,
        caps=caps,
        can_upload_ui=can_upload_ui,
        cannot_upload_reason=deny_reason,
        user_display_name=user_display_name,
    )
